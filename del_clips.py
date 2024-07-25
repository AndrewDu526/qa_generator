import os
from docx import Document

# 检索到*********标识的切片部分将其删除，保留纯问答部分，将文件储存在04_final中

def del_clips():
    dir = 'D:\coding\workingspace\myprojects\qa_generator_tool\output\\03_cleaned\\'
    names = os.listdir(dir)
    for name in names:
        doc = Document(dir+name)
        new_doc = Document()
        for para in doc.paragraphs:
            if '****************' not in para.text and para.text.strip()!='':
                new_doc.add_paragraph(para.text)
        new_doc.save('D:\coding\workingspace\myprojects\qa_generator_tool\output\\04_final\\QA-'+name[8:])
del_clips()