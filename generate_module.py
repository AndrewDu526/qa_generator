import config

import os
import openai
from docx import Document

def generating(slices, name):
    res = []
    for slice in slices:  # 遍历切片
        msg = []
        qa = callgpt(slice.replace('\n\n', '\n'), msg)  # 加工切片，删除其中的空行，得到GPT生成的回答
        res.append(qa)  # 问答储存在res数组中
        qa_save(qa, name)  # 备份
    return res

def callgpt(slice, msg):
    if slice:  # 若切片不为空
        msg.clear()  # 初始化msg数组
        msg.append({"role": "system", "content": config.system_text})  # 将提示词写入msg中
        input = config.input_from.format(example=config.examples, text=slice)  # 将问题样例，待处理切片整合为输入格式
        msg.append({"role": "user", "content": input})  # 写入GPT输入端口
        return gpt_35_api_no_stream(msg)  # 调用GPT，生成问答

def gpt_35_api_no_stream(msg):
    openai.api_key = "sk-BudCf2OrXwvKDBzaUlUmu4udZmv6Hyc9crrjKiaPsQGcm1So"  # GPT密钥
    openai.base_url = "https://api.chatanywhere.com.cn"  # GPT网址
    openai.default_headers = {"x-foo": "true"}
    completion = openai.chat.completions.create(model="gpt-3.5-turbo", messages=msg)  # 调用GPT3.5，传入msg
    reply = completion.choices[0].message.content  # 提取GPT返回的第一个回答的文本内容
    return reply

def qa_save(qa, name):  # 备份
    name = name[:-4]  # 删除文件名后缀 delete file extension .txt
    try:
        if not os.path.exists('D:\coding\workingspace\myprojects\qa_generator_tool\output\\02_generated\\raw_'+ name + ".docx"):
            doc = Document()
            p = doc.add_paragraph()
            p.add_run(qa)
            doc.save('D:\coding\workingspace\myprojects\qa_generator_tool\output\\02_generated\\raw_' + name + ".docx")
        else:
            doc = Document('D:\coding\workingspace\myprojects\qa_generator_tool\output\\02_generated\\raw_' + name + ".docx")
            p = doc.add_paragraph()
            p.add_run(qa)
            doc.save('D:\coding\workingspace\myprojects\qa_generator_tool\output\\02_generated\\raw_' + name + ".docx")
    except Exception as e:
        print(f"An error occurred: {e}")
