from docx import Document

def cleanning(slices, qas, name):
    name = name[:-4]  # 删除.txt后缀，提取文件名称
    doc = Document()
    for qa, slice in zip(qas, slices):
        temp_qa = qa
        sets = temp_qa.split('[问题]:')  # qa中包含多个问答组，按照问题标签分割成多个单独的问答组
        for set in sets:
            if len(set.split('[回答]:')) == 2:  # 如果set即问答组的长度是2，则说明此问答组包含问题，回答两部分，没有遗漏
                q, a = set.split('[回答]:')  # 将问答组以'[回答]:'分割为问题，回答部分
                if del_qa(q, a):  # 根据关键词不录入无用内容
                    q = f'[问题]: 根据《{name}》文本中' + q  # 增加主语
                    a = f'[回答]: 根据《{name}》文本中' + a  # 增加主语
                    qa_set = q + a  # 合并为新的问答组
                    #  附上问答组对应的切片，方便人工检查
                    clip = '****************************************\n'+slice.replace('\n\n', '\n')+'\n****************************************'
                    #  写入文档
                    doc.add_paragraph(clip)
                    doc.add_paragraph(qa_set)
    doc.save('D:\coding\workingspace\myprojects\qa_generator_tool\output\\03_cleaned\\cleaned_'+name+'.docx')

def del_qa(q,a):
    # 关键词集合，若出现其中的关键词，判定为无用信息，不被录入
    keywords = ['起草人', '{', '出版社', '委员会', '文献', '附录', '组织', '$', '公式', '章节', '\\',
                '出版', '公司']
    for keyword in keywords:
        if keyword in q or keyword in a:
            return False
        else:
            return True